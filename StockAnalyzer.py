# Imports
import yfinance as yf
from datetime import datetime
from os import path
from docx import Document


class Stock:
    def __init__(self, symbol):
        self.data = yf.Ticker(symbol)
    
    def extract_info(self):
        """
        Extracts intresting info about the stock
        """

        self.name = self.extract_value("longName")
        self.symbol = self.extract_value("symbol")
        self.business_summary = self.extract_value("longBusinessSummary")
        self.industry = self.extract_value("industry")
        self.sector = self.extract_value("sector")
        self.country = self.extract_value("country")
        self.current_price = self.extract_value("currentPrice")
        self.market_cap = self.extract_value("marketCap")
        self.volume = self.extract_value("volume")
        self.average_volume = self.extract_value("averageVolume")
        self.trailing_pe = self.extract_value("trailingPE")
        self.forward_pe = self.extract_value("forwardPE")
        self.beta = self.extract_value("beta")
        self.dividend_yield = self.extract_value("yield")
        self.fifty_two_week_high = self.extract_value("fiftyTwoWeekHigh")
        self.fifty_two_week_low = self.extract_value("fiftyTwoWeekLow")
        self.two_hundred_day_average = self.extract_value("twoHundredDayAverage")
        self.fifty_two_week_change = self.extract_value("52WeekChange")
        self.s_and_p_fifty_two_week_change = self.extract_value("SandP52WeekChange")
    
    def extract_value(self, key):
        try:
            value = self.data.info[key]
            if value:
                return value
            return "None"
        except KeyError:
            return "None"
    
    def analyze_earnings(self):

        # Extracts stock's last four years earnings
        self.earnings = self.data.get_earnings().to_dict()["Earnings"]
        current_earnings = list(self.earnings.values())[-1]
        for year, earnings in list(self.earnings.items())[:-1]:
            if (earnings > 0):
                growth = (current_earnings - earnings) / earnings
            else:
                growth = (current_earnings - earnings) / earnings * -1
            print(f"Earning's growth since {year}: {growth:%}")
    
    def write_report(self):
        
        # Example for report name: aapl(11-04-22).txt
        self.report_name = f"{self.symbol}({datetime.now().strftime('%d-%m-%y')}).docx"
        
        # Initializes word document object
        self.document = Document()

        # Setting the docunebt's header as title
        self.document.add_heading(f'{self.name} ({self.symbol})', 0)

        # Setting the general information header as level 1
        self.document.add_heading('General Information', level=1)

        # Business summary paragraph
        self.write_paragraph("Business Summary", self.business_summary)
        
        # Industry paragraph
        self.write_paragraph("Industry", self.industry)

        # Sector paragraph
        self.write_paragraph("Sector", self.sector)

        # Country paragraph
        self.write_paragraph("Country", self.country)

        # Setting the technical information header as level 1
        self.document.add_heading('Technical Information', level=1)

        # Current price paragraph
        self.write_paragraph("Current Price", self.current_price)

        # Market cap paragraph
        self.write_paragraph("Market Cap", self.market_cap)

        # Volume paragraph
        self.write_paragraph("Volume", self.volume)

        # Average volume paragraph
        self.write_paragraph("Average Volume", self.average_volume)

        # Trailing pe ratio paragraph
        self.write_paragraph("Trailing PE Ratio", self.trailing_pe)

        # Forward pe ratio paragraph
        self.write_paragraph("Forward PE Ratio", self.forward_pe)

        # Beta paragraph
        self.write_paragraph("Beta", self.beta)

        # Dividend yield paragraph
        self.write_paragraph("Dividend Yield", self.dividend_yield, percentage=True)

        # S&P 500 52 week change paragraph
        self.write_paragraph("S&P 500's 52 Week Change", self.s_and_p_fifty_two_week_change, percentage=True)

        # 52 week change paragraph
        self.write_paragraph("52 Week Change", self.fifty_two_week_change, percentage=True)

        # 52 week high paragraph
        self.write_paragraph("52 Week High", self.fifty_two_week_high)

        # 52 week low paragraph
        self.write_paragraph("52 Week Low", self.fifty_two_week_low)

        # 200 days average paragraph
        self.write_paragraph("200 Days Average", self.two_hundred_day_average)

        # Saves the report
        self.document.save(self.report_name)

    def write_paragraph(self, key, value, percentage=False):

        # Checks whether the value is an int and bigger than 1000 and formats like this 1,000
        if (type(value) == int) and (value > 1000) and not percentage:
            value = f"{value:,}"
        
        # Checks whether value is float and formats to have to digits after the decimal point
        if (type(value) == float) and not percentage:
            value = f"{value:.2f}"
        
        if percentage and (value != "None"):
            value = f"{value:.2%}"
        
        # Creating new paragraph
        p = self.document.add_paragraph()

        # Adding the subject with underline
        p.add_run(f"{key}:").underline = True

        # Adding the value
        p.add_run(f'\n{value}')



def main():
    #stock_symbol = input("Enter stock symbol: ")
    stock_symbol = "aapl"
    my_stock = Stock(stock_symbol)
    my_stock.extract_info()
    #print(f"{my_stock.dividend_yield:.2%}")
    my_stock.write_report()


if __name__ == "__main__":
    main()
